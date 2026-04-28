import os
import re
import io
import time
import hashlib
import pandas as pd
import plotly.express as px
import streamlit as st
from dotenv import load_dotenv

from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

try:
    import numpy as np
    HAS_NUMPY = True
except ImportError:
    HAS_NUMPY = False

try:
    from sklearn.linear_model import LinearRegression
    from sklearn.preprocessing import PolynomialFeatures
    from sklearn.metrics import r2_score
    HAS_SKLEARN = True
except ImportError:
    HAS_SKLEARN = False

try:
    from prophet import Prophet
    HAS_PROPHET = True
except ImportError:
    HAS_PROPHET = False

try:
    import sqlalchemy
    from sqlalchemy import create_engine, text, inspect
    HAS_SQLALCHEMY = True
except ImportError:
    HAS_SQLALCHEMY = False

try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Load .env on startup
load_dotenv()


# ─────────────────────────────────────────────────────────────
# MODEL REGISTRY
# ─────────────────────────────────────────────────────────────

MODELS = {
    "gemini-2.5-flash": {
        "label": "Gemini 2.5 Flash",
        "provider": "google",
        "icon": "🟢",
        "color": "#4ade80",
        "description": "Google AI Studio · Frontier quality · 1M context",
    },
    "gpt-4.1": {
        "label": "GPT-4.1",
        "provider": "github",
        "icon": "🤖",
        "color": "#7c9ff5",
        "description": "GitHub Models · Best quality · 1M context · Instruction following",
        "model_name": "openai/gpt-4.1",
    },
    "gpt-4.1-mini": {
        "label": "GPT-4.1 Mini",
        "provider": "github",
        "icon": "⚡",
        "color": "#f59e0b",
        "description": "GitHub Models · Faster & lighter · 1M context · Lower latency",
        "model_name": "openai/gpt-4.1-mini",
    },
    "gpt-4.1-nano": {
        "label": "GPT-4.1 Nano",
        "provider": "github",
        "icon": "🚀",
        "color": "#c084fc",
        "description": "GitHub Models · Lowest latency · 1M context · Simple tasks",
        "model_name": "openai/gpt-4.1-nano",
    },
}

# Default to Gemini since that key is present in .env
DEFAULT_MODEL = os.getenv("DEFAULT_MODEL", "gemini-2.5-flash")

MAX_DISPLAY_ROWS = 500
MAX_HISTORY_TURNS = 6   # last N user+assistant turns kept for context
FORECAST_HORIZON_DEFAULT = 10   # default periods to forecast ahead

# Fallback chain: when a model is rate-limited, try the next one in this list
FALLBACK_CHAIN = ["gemini-2.5-flash", "gpt-4.1-mini", "gpt-4.1-nano", "gpt-4.1"]


# ─────────────────────────────────────────────────────────────
# RATE-LIMIT DETECTION
# ─────────────────────────────────────────────────────────────

class RateLimitError(Exception):
    """Raised when a model returns a 429 / quota-exceeded response."""
    pass


def _is_rate_limit(exc: Exception) -> bool:
    """Detect 429 / quota errors across providers."""
    msg = str(exc).lower()
    return any(k in msg for k in [
        "429", "quota", "rate limit", "rate_limit", "resource_exhausted",
        "too many requests", "exceeded your current quota",
    ])


# ─────────────────────────────────────────────────────────────
# LLM CLIENT FACTORY
# ─────────────────────────────────────────────────────────────

def build_llm_client(model_id: str):
    """
    Returns a callable: complete(system, user, max_tokens, history=None) -> str
    history is a list of {"role": "user"|"assistant", "content": str}
    """
    cfg = MODELS[model_id]
    provider = cfg["provider"]

    if provider == "github":
        # GitHub Models uses OpenAI-compatible API
        from openai import OpenAI
        token = os.getenv("GITHUB_TOKEN") or os.getenv("OPENAI_API_KEY", "")
        if not token:
            raise ValueError(
                "GitHub Models requires GITHUB_TOKEN (or OPENAI_API_KEY) in your .env. "
                "Get a free token at github.com/settings/tokens"
            )
        client = OpenAI(
            base_url="https://models.github.ai/inference",
            api_key=token,
        )
        model_name = cfg["model_name"]

        def complete(system: str, user: str, max_tokens: int = 1024,
                     history: list | None = None) -> str:
            messages = []
            if system:
                messages.append({"role": "system", "content": system})
            for h in (history or []):
                messages.append({"role": h["role"], "content": h["content"]})
            messages.append({"role": "user", "content": user})
            try:
                resp = client.chat.completions.create(
                    model=model_name,
                    messages=messages,
                    max_tokens=max_tokens,
                    temperature=0.0,
                )
                return resp.choices[0].message.content.strip()
            except Exception as e:
                if _is_rate_limit(e):
                    raise RateLimitError(str(e)) from e
                raise

        return complete

    elif provider == "google":
        import google.generativeai as genai
        api_key = os.getenv("GEMINI_API_KEY", "")
        if not api_key:
            raise ValueError(
                "Gemini requires GEMINI_API_KEY in your .env. "
                "Get a free key at aistudio.google.com/apikey"
            )
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(
            model_name="gemini-2.5-flash",
            system_instruction="You are a helpful, concise assistant.",
        )

        def complete(system: str, user: str, max_tokens: int = 1024,
                     history: list | None = None) -> str:
            parts = []
            if system:
                parts.append(f"[System]: {system}")
            for h in (history or []):
                role_label = "User" if h["role"] == "user" else "Assistant"
                parts.append(f"[{role_label}]: {h['content']}")
            parts.append(f"[User]: {user}")
            full_prompt = "\n\n".join(parts)
            try:
                resp = model.generate_content(
                    full_prompt,
                    generation_config=genai.GenerationConfig(
                        max_output_tokens=max_tokens,
                        temperature=0.0,
                    ),
                )
                return resp.text.strip()
            except Exception as e:
                if _is_rate_limit(e):
                    raise RateLimitError(str(e)) from e
                raise

        return complete

    else:
        raise ValueError(f"Unknown provider: {provider}")


@st.cache_resource(show_spinner="⚡ Loading model...")
def get_llm_client(model_id: str):
    """Cached per model_id."""
    return build_llm_client(model_id)


def _available_fallbacks(current_model_id: str) -> list[str]:
    """Return fallback models that have credentials configured, excluding current."""
    has_gemini = bool(os.getenv("GEMINI_API_KEY"))
    has_github = bool(os.getenv("GITHUB_TOKEN") or os.getenv("OPENAI_API_KEY"))
    available = []
    for mid in FALLBACK_CHAIN:
        if mid == current_model_id:
            continue
        cfg = MODELS.get(mid, {})
        if cfg.get("provider") == "google" and not has_gemini:
            continue
        if cfg.get("provider") == "github" and not has_github:
            continue
        available.append(mid)
    return available


# Per-model client cache (plain dict, not st.cache_resource so exceptions stay catchable)
_client_cache: dict = {}

def _get_client(model_id: str):
    if model_id not in _client_cache:
        _client_cache[model_id] = build_llm_client(model_id)
    return _client_cache[model_id]


def make_fallback_complete(primary_model_id: str):
    """
    Returns a complete() function that automatically falls back through
    FALLBACK_CHAIN when the primary model hits a rate limit / quota error.
    """
    def complete_with_fallback(system: str, user: str, max_tokens: int = 1024,
                               history: list | None = None) -> str:
        chain = [primary_model_id] + _available_fallbacks(primary_model_id)
        last_exc = None
        for model_id in chain:
            try:
                fn = _get_client(model_id)
                result = fn(system, user, max_tokens, history)
                if model_id != primary_model_id:
                    st.session_state["_last_fallback"] = model_id
                    st.toast(
                        f"✅ Fell back to {MODELS[model_id]['icon']} {MODELS[model_id]['label']}",
                        icon="🔄",
                    )
                else:
                    st.session_state.pop("_last_fallback", None)
                return result
            except RateLimitError as e:
                last_exc = e
                fallbacks_left = chain[chain.index(model_id)+1:]
                if fallbacks_left:
                    next_m = MODELS.get(fallbacks_left[0], {})
                    st.toast(
                        f"⚠️ {MODELS[model_id]['label']} rate-limited — "
                        f"switching to {next_m.get('icon','')} {next_m.get('label', fallbacks_left[0])}",
                        icon="🔄",
                    )
                continue
            except Exception:
                raise
        # All models exhausted — show a friendly error instead of crashing
        st.error(f"⚠️ All models are currently rate-limited. Please wait a moment and try again. Last error: {last_exc}")
        st.stop()
    return complete_with_fallback


# ─────────────────────────────────────────────────────────────
# PAGE CONFIG & STYLES
# ─────────────────────────────────────────────────────────────

st.set_page_config(page_title="DocBot", page_icon="🤖", layout="wide")

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;600&family=IBM+Plex+Sans:wght@400;500;600;700&display=swap');

    * { font-family: 'IBM Plex Sans', sans-serif; }
    code, pre, .badge-mono { font-family: 'IBM Plex Mono', monospace; }

    .stApp { background-color: #0a0d14; color: #e0e0e0; }
    [data-testid="stSidebar"] {
        background-color: #0f1420;
        border-right: 1px solid #1e2535;
    }
    [data-testid="stChatMessage"] {
        background-color: #111827;
        border: 1px solid #1e2535;
        border-radius: 12px;
        margin-bottom: 8px;
        padding: 4px 0;
    }
    [data-testid="stChatInput"] textarea {
        background-color: #111827 !important;
        color: #e0e0e0 !important;
        border: 1px solid #2d3748 !important;
        border-radius: 10px !important;
        font-family: 'IBM Plex Sans', sans-serif !important;
    }
    h1, h2 { color: #818cf8 !important; font-weight: 700 !important; letter-spacing: -0.5px; }
    h3 { color: #a5b4fc !important; }

    .badge-data  { background:#1a2e4a; color:#60a5fa;  padding:3px 12px; border-radius:20px; font-size:11px; font-weight:700; letter-spacing:0.5px; text-transform:uppercase; }
    .badge-rag   { background:#1a3a2a; color:#4ade80;  padding:3px 12px; border-radius:20px; font-size:11px; font-weight:700; letter-spacing:0.5px; text-transform:uppercase; }
    .badge-graph { background:#2d1a4a; color:#c084fc;  padding:3px 12px; border-radius:20px; font-size:11px; font-weight:700; letter-spacing:0.5px; text-transform:uppercase; }
    .badge-error { background:#3a1a1a; color:#f87171;  padding:3px 12px; border-radius:20px; font-size:11px; font-weight:700; letter-spacing:0.5px; text-transform:uppercase; }
    .badge-forecast { background:#2a1e10; color:#f59e0b; padding:3px 12px; border-radius:20px; font-size:11px; font-weight:700; letter-spacing:0.5px; text-transform:uppercase; }
    .badge-sql   { background:#1a2a3a; color:#38bdf8;  padding:3px 12px; border-radius:20px; font-size:11px; font-weight:700; letter-spacing:0.5px; text-transform:uppercase; }

    .sql-conn-card { background:#0d1829; border:1px solid #1e3a5f; border-radius:12px; padding:16px 18px; margin:8px 0; }
    .sql-table-chip { background:#0f2237; border:1px solid #1e3a5f; color:#38bdf8; border-radius:6px; padding:3px 10px; font-size:11px; display:inline-block; margin:3px; font-family:'IBM Plex Mono',monospace; cursor:default; }
    .sql-connected { background:#0d1f18; border:1px solid #134e35; border-radius:8px; padding:8px 12px; color:#4ade80; font-size:12px; font-family:'IBM Plex Mono',monospace; }

    .info-box    { background:#111827; border:1px solid #1e2535; border-left:3px solid #818cf8; border-radius:8px; padding:12px 16px; margin:8px 0; font-size:13px; color:#9ca3af; line-height:1.7; }
    .upload-hint { background:#0f1420; border:2px dashed #1e2535; border-radius:16px; padding:32px; text-align:center; color:#6b7280; font-size:14px; margin:20px 0; }
    .model-card  { border-radius:8px; padding:10px 14px; margin:4px 0; font-size:12px; border-width:1px; border-style:solid; line-height:1.6; }
    .stat-card   { background:#111827; border:1px solid #1e2535; border-radius:8px; padding:12px 16px; text-align:center; }
    .timing-note { color:#4b5563; font-size:11px; font-family:'IBM Plex Mono',monospace; margin-top:4px; }
    .source-chip { background:#1a1f2e; border:1px solid #2d3748; color:#9ca3af; border-radius:6px; padding:2px 8px; font-size:11px; display:inline-block; margin:2px; font-family:'IBM Plex Mono',monospace; }
    .turn-counter { color:#4b5563; font-size:11px; margin-top:4px; }

    div[data-testid="stSelectbox"] label { color: #9ca3af !important; font-size: 13px !important; }
    div[data-testid="stButton"] button {
        border-radius: 20px !important;
        font-size: 12px !important;
        padding: 4px 12px !important;
        transition: all 0.15s;
    }
    /* Download button */
    div[data-testid="stDownloadButton"] button {
        border-radius: 8px !important;
        font-size: 12px !important;
        background: #1a2e4a !important;
        color: #60a5fa !important;
        border: 1px solid #2d4a6a !important;
    }
    .stAlert { border-radius: 10px !important; }

    /* Scrollable dataframe */
    [data-testid="stDataFrame"] { border-radius: 10px; overflow: hidden; }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────
# FILE PARSERS
# ─────────────────────────────────────────────────────────────

def parse_pdf(file_bytes: bytes, filename: str) -> list[Document]:
    if not HAS_PYPDF:
        st.warning("pypdf not installed — cannot parse PDFs.")
        return []
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": filename, "page": i + 1, "type": "pdf"}
            ))
    return docs


def parse_docx(file_bytes: bytes, filename: str) -> list[Document]:
    if not HAS_DOCX:
        st.warning("python-docx not installed — cannot parse Word docs.")
        return []
    doc = DocxDocument(io.BytesIO(file_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if not full_text.strip():
        return []
    return [Document(page_content=full_text, metadata={"source": filename, "type": "docx"})]


def parse_txt(file_bytes: bytes, filename: str) -> list[Document]:
    text = file_bytes.decode("utf-8", errors="ignore")
    if not text.strip():
        return []
    return [Document(page_content=text, metadata={"source": filename, "type": "txt"})]


def parse_csv(file_bytes: bytes, filename: str) -> tuple[list[Document], pd.DataFrame | None]:
    try:
        df = pd.read_csv(io.BytesIO(file_bytes))
    except Exception as e:
        st.warning(f"Could not parse CSV {filename}: {e}")
        return [], None
    schema_text = (
        f"File: {filename}\nColumns: {list(df.columns)}\nShape: {df.shape}\n"
        f"Data types: {df.dtypes.to_dict()}\n\n"
        f"Sample (first 5 rows):\n{df.head(5).to_string(index=False)}"
    )
    stats_text = f"Statistics for {filename}:\n{df.describe(include='all').to_string()}"
    return [
        Document(page_content=schema_text, metadata={"source": filename, "type": "csv_schema"}),
        Document(page_content=stats_text,  metadata={"source": filename, "type": "csv_stats"}),
    ], df


def parse_excel(file_bytes: bytes, filename: str) -> tuple[list[Document], pd.DataFrame | None]:
    try:
        df = pd.read_excel(io.BytesIO(file_bytes))
    except Exception as e:
        st.warning(f"Could not parse Excel {filename}: {e}")
        return [], None
    schema_text = (
        f"File: {filename}\nColumns: {list(df.columns)}\nShape: {df.shape}\n"
        f"Data types: {df.dtypes.to_dict()}\n\n"
        f"Sample (first 5 rows):\n{df.head(5).to_string(index=False)}"
    )
    stats_text = f"Statistics for {filename}:\n{df.describe(include='all').to_string()}"
    return [
        Document(page_content=schema_text, metadata={"source": filename, "type": "excel_schema"}),
        Document(page_content=stats_text,  metadata={"source": filename, "type": "excel_stats"}),
    ], df


# ─────────────────────────────────────────────────────────────
# VECTORSTORE BUILDER  (cache keyed on stable content hash)
# ─────────────────────────────────────────────────────────────

def _content_hash(file_contents: tuple) -> str:
    h = hashlib.md5()
    for name, data in file_contents:
        h.update(name.encode())
        h.update(data)
    return h.hexdigest()


@st.cache_resource(show_spinner="🧠 Building knowledge base...")
def build_vectorstore(content_hash: str, file_contents: tuple) -> tuple:  # noqa: ARG001
    """content_hash is used as the cache key; file_contents provides data."""
    all_docs: list[Document] = []
    dataframes: dict[str, pd.DataFrame] = {}

    for filename, file_bytes in file_contents:
        ext = filename.rsplit(".", 1)[-1].lower()
        if ext == "pdf":
            all_docs.extend(parse_pdf(file_bytes, filename))
        elif ext == "docx":
            all_docs.extend(parse_docx(file_bytes, filename))
        elif ext == "txt":
            all_docs.extend(parse_txt(file_bytes, filename))
        elif ext == "csv":
            docs, df = parse_csv(file_bytes, filename)
            all_docs.extend(docs)
            if df is not None:
                dataframes[filename] = df
        elif ext in ("xlsx", "xls"):
            docs, df = parse_excel(file_bytes, filename)
            all_docs.extend(docs)
            if df is not None:
                dataframes[filename] = df

    if not all_docs:
        return None, dataframes

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
    splits = splitter.split_documents(all_docs)
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    vectorstore = FAISS.from_documents(splits, embeddings)
    return vectorstore, dataframes


# ─────────────────────────────────────────────────────────────
# INTENT CLASSIFIER  (robust multi-label extraction)
# ─────────────────────────────────────────────────────────────

def classify_intents(complete, question: str, has_tabular: bool, has_sql: bool = False) -> list[str]:
    """
    Returns an ORDERED list of intents to execute, e.g. ["RAG", "GRAPH"].
    GRAPH is allowed even without tabular data — we extract numbers from text docs.
    FORECAST is for predictive analytics, trend lines, and future value prediction.
    SQL is for querying a connected database with natural language.
    """
    # Fast-path: keyword override for unambiguous forecast requests
    _q_lower = question.lower()
    _forecast_keywords = [
        "forecast", "predict", "prediction", "extrapolate", "extrapolation",
        "trend line", "trendline", "linear regression", "polynomial regression",
        "next month", "next year", "next quarter", "next week", "next period",
        "future value", "future sales", "future revenue", "time series",
        "regress", "regression",
    ]
    _sql_keywords = [
        "sql", "database", "db", "table", "select", "query the", "from the db",
        "from database", "from the database", "join", "where clause",
        "run a query", "write a query", "fetch from", "records in",
    ]
    _has_forecast_kw = any(kw in _q_lower for kw in _forecast_keywords)
    _has_sql_kw = has_sql and any(kw in _q_lower for kw in _sql_keywords)

    sql_label_desc = (
        "  SQL     = query the connected SQL database with natural language\n"
        if has_sql else ""
    )
    sql_examples = (
        "  'show me all users from the database' → SQL\n"
        "  'how many orders were placed last month in the db' → SQL\n"
        "  'join orders and customers and show revenue by region' → SQL\n"
        if has_sql else ""
    )

    system = (
        "You are an intent classifier for a document assistant. "
        "A user message may request ONE or MULTIPLE actions simultaneously.\n\n"
        "Labels:\n"
        "  RAG      = summarize, explain, find info, Q&A from text/PDF documents\n"
        "  GRAPH    = ANY request for a chart, plot, graph, visualization — including on PDF/text data\n"
        "  DATA     = filter rows, sort, top-N from a spreadsheet or CSV file\n"
        "  FORECAST = predict future values, forecast, trend analysis, regression, time-series, "
        "             'what will X be next month', 'predict', 'extrapolate', 'trend line', "
        "             'linear regression', 'next N periods'\n"
        + sql_label_desc +
        "\nIMPORTANT: If the user asks to predict, forecast, extrapolate, or show a trend line, "
        "ALWAYS include FORECAST. Never substitute RAG for FORECAST.\n\n"
        "Reply with a comma-separated list of applicable labels IN EXECUTION ORDER.\n"
        "Examples:\n"
        "  'summarize the report' → RAG\n"
        "  'make a chart' → GRAPH\n"
        "  'predict next 5 months of sales' → FORECAST\n"
        "  'forecast revenue and show a chart' → FORECAST,GRAPH\n"
        "  'what will the trend look like' → FORECAST\n"
        "  'run a linear regression on the data' → FORECAST\n"
        "  'predict future values and show trend line' → FORECAST\n"
        "  'forecast the next 10 periods using linear regression' → FORECAST\n"
        "  'summarize and make a chart' → RAG,GRAPH\n"
        "  'show top 10 rows and plot them' → DATA,GRAPH\n"
        "  'what does it say and show a chart' → RAG,GRAPH\n"
        + sql_examples +
        "\nReply with ONLY the labels, comma-separated, no explanation."
    )
    try:
        raw = complete(system, question, 30)
    except Exception:
        if _has_sql_kw:
            return ["SQL"]
        return ["FORECAST"] if _has_forecast_kw else ["RAG"]

    upper = raw.strip().upper()
    positions = {}
    for label in ["RAG", "GRAPH", "DATA", "FORECAST", "SQL"]:
        idx = upper.find(label)
        if idx != -1:
            positions[label] = idx
    ordered = sorted(positions.keys(), key=lambda l: positions[l])
    result = ordered if ordered else ["RAG"]

    # Override: keyword-based injection
    if _has_forecast_kw and "FORECAST" not in result:
        if result == ["RAG"]:
            result = ["FORECAST"]
        else:
            result = ["FORECAST"] + [r for r in result if r != "RAG"]
    if _has_sql_kw and "SQL" not in result:
        result = ["SQL"] + [r for r in result if r != "SQL"]

    # If SQL intent but no DB connected, drop it silently (pipeline will warn)
    if "SQL" in result and not has_sql:
        result = [r for r in result if r != "SQL"]

    return result if result else ["RAG"]


# ─────────────────────────────────────────────────────────────
# PIPELINE A — GRAPH
# ─────────────────────────────────────────────────────────────

def run_graph_from_context(complete, vectorstore, question: str, history: list):
    """Extract numeric data from text/PDF context, build a DataFrame, then plot."""
    if vectorstore is None:
        return None, "No documents loaded."
    docs = vectorstore.as_retriever(search_kwargs={"k": 10}).invoke(question)
    if not docs:
        return None, "No relevant content found in documents."
    context = "\n\n---\n\n".join(
        f"[Source: {d.metadata.get('source', 'unknown')}]\n{d.page_content}"
        for d in docs
    )
    # Step 1: extract structured data as JSON
    extract_system = (
        "You are a data extraction expert. Given the document context below, "
        "extract ALL numeric/tabular data relevant to the user's chart request.\n"
        "Return a JSON object with:\n"
        "  \"data\": list of row objects (e.g. [{\"label\": \"A\", \"value\": 10}, ...])\n"
        "  \"x_col\": name of the category/label column\n"
        "  \"y_col\": name of the numeric column\n"
        "  \"title\": a descriptive chart title\n"
        "  \"chart_type\": one of: bar, line, pie, scatter, histogram\n\n"
        "Return ONLY valid JSON, no markdown fences, no explanation.\n\n"
        f"Context:\n{context}"
    )
    recent_hist = [h for h in history[-MAX_HISTORY_TURNS*2:]
                   if h["role"] in ("user", "assistant") and h.get("type") == "text"]
    try:
        raw_json = complete(extract_system, question, 1000, history=recent_hist)
        # Strip possible markdown fences
        raw_json = re.sub(r"^```[a-z]*\n?", "", raw_json.strip(), flags=re.MULTILINE)
        raw_json = re.sub(r"```$", "", raw_json.strip(), flags=re.MULTILINE)
        import json
        parsed = json.loads(raw_json.strip())
        data = parsed["data"]
        x_col = parsed.get("x_col", "label")
        y_col = parsed.get("y_col", "value")
        title = parsed.get("title", "Chart")
        chart_type = parsed.get("chart_type", "bar")
    except Exception as e:
        return None, f"Could not extract chart data from document: {e}\n\nRaw response:\n{raw_json if 'raw_json' in dir() else '(no response)'}"

    try:
        df = pd.DataFrame(data)
        if x_col not in df.columns or y_col not in df.columns:
            # Try to infer columns
            cols = df.columns.tolist()
            x_col = cols[0]
            y_col = next((c for c in cols[1:] if pd.api.types.is_numeric_dtype(df[c])), cols[-1])
        df[y_col] = pd.to_numeric(df[y_col], errors="coerce")
        df = df.dropna(subset=[y_col])
    except Exception as e:
        return None, f"Could not build DataFrame from extracted data: {e}"

    try:
        chart_fn = {
            "bar": px.bar, "line": px.line, "pie": px.pie,
            "scatter": px.scatter, "histogram": px.histogram,
        }.get(chart_type, px.bar)
        if chart_type == "pie":
            fig = chart_fn(df, names=x_col, values=y_col, title=title)
        else:
            fig = chart_fn(df, x=x_col, y=y_col, title=title)
        fig.update_layout(
            paper_bgcolor="#0a0d14", plot_bgcolor="#111827",
            font_color="#e0e0e0", font_family="IBM Plex Sans",
            title_font_color="#818cf8",
            legend=dict(bgcolor="#111827", bordercolor="#1e2535"),
            xaxis=dict(gridcolor="#1e2535", zeroline=False),
            yaxis=dict(gridcolor="#1e2535", zeroline=False),
        )
        return fig, None
    except Exception as e:
        return None, f"Chart rendering error: {e}"


def run_graph_pipeline(complete, dataframes: dict, question: str, history: list):
    if not dataframes:
        return None, "NO_TABULAR"  # signal caller to use run_graph_from_context instead

    df_summary = "\n".join(
        f"DataFrame '{name}': columns={list(df.columns)}, shape={df.shape}, "
        f"dtypes={df.dtypes.to_dict()}"
        for name, df in dataframes.items()
    )
    system = (
        "You are a Python data visualisation expert. Available DataFrames (already in scope):\n"
        f"{df_summary}\n\n"
        "Write Python code to create a Plotly Express chart assigned to variable `fig`.\n"
        "Use ONLY: `df` (default first dataframe), `px`, `pd`, and named df vars.\n"
        "Handle missing values with dropna() where needed.\n"
        "Return ONLY a fenced python code block, nothing else:\n```python\nfig = ...\n```"
    )
    recent_hist = [h for h in history[-MAX_HISTORY_TURNS*2:] if h["role"] in ("user", "assistant") and h.get("type") == "text"]
    try:
        response = complete(system, question, 600, history=recent_hist)
    except Exception as e:
        return None, f"LLM error: {e}"

    match = re.search(r'```python\n(.*?)```', response, re.DOTALL)
    code = match.group(1).strip() if match else response.strip()
    # Strip any import lines — only px/pd/df are available
    code_lines = [l for l in code.splitlines() if not l.strip().startswith("import")]
    code = "\n".join(code_lines)

    default_df = next(iter(dataframes.values()))
    local_vars = {"df": default_df, "px": px, "pd": pd}
    for name, frame in dataframes.items():
        safe = re.sub(r'[^a-zA-Z0-9_]', '_', name.rsplit('.', 1)[0])
        local_vars[safe] = frame

    try:
        exec(code, {"__builtins__": {}}, local_vars)  # restricted builtins
        fig = local_vars.get("fig")
        if fig is None:
            return None, f"No `fig` variable created.\n\n```python\n{code}\n```"
        fig.update_layout(
            paper_bgcolor="#0a0d14", plot_bgcolor="#111827",
            font_color="#e0e0e0", font_family="IBM Plex Sans",
            title_font_color="#818cf8",
            legend=dict(bgcolor="#111827", bordercolor="#1e2535"),
            xaxis=dict(gridcolor="#1e2535", zeroline=False),
            yaxis=dict(gridcolor="#1e2535", zeroline=False),
        )
        return fig, None
    except Exception as e:
        # Try once more with full builtins in case restricted eval failed
        try:
            exec(code, {}, local_vars)
            fig = local_vars.get("fig")
            if fig:
                fig.update_layout(
                    paper_bgcolor="#0a0d14", plot_bgcolor="#111827",
                    font_color="#e0e0e0", font_family="IBM Plex Sans",
                    title_font_color="#818cf8",
                )
                return fig, None
        except Exception:
            pass
        return None, f"Chart error: {e}\n\n```python\n{code}\n```"


# ─────────────────────────────────────────────────────────────
# PIPELINE B — DATA  (now includes NL summary)
# ─────────────────────────────────────────────────────────────

def run_data_pipeline(complete, dataframes: dict, question: str, history: list) -> dict:
    if not dataframes:
        return {"code": "", "df": None, "summary": None,
                "error": "No tabular data available. Upload a CSV or Excel file.", "capped": False}

    df_summary = "\n".join(
        f"DataFrame '{name}': columns={list(df.columns)}, shape={df.shape}\n"
        f"  dtypes: {df.dtypes.to_dict()}\n"
        f"  Sample: {df.head(2).to_dict('records')}"
        for name, df in dataframes.items()
    )
    system = (
        "You are a pandas expert. Available DataFrames:\n"
        f"{df_summary}\n\n"
        "Write Python to produce a DataFrame or scalar assigned to `result`.\n"
        "Use ONLY `df`, `pd`, and the named df variables.\n"
        "Handle NaN values gracefully. Do NOT import anything.\n"
        "Return ONLY:\n```python\nresult = ...\n```"
    )
    recent_hist = [h for h in history[-MAX_HISTORY_TURNS*2:] if h["role"] in ("user", "assistant") and h.get("type") == "text"]
    try:
        response = complete(system, question, 512, history=recent_hist)
    except Exception as e:
        return {"code": "", "df": None, "summary": None, "error": f"LLM error: {e}", "capped": False}

    match = re.search(r'```python\n(.*?)```', response, re.DOTALL)
    code = match.group(1).strip() if match else response.strip()
    code_lines = [l for l in code.splitlines() if not l.strip().startswith("import")]
    code = "\n".join(code_lines)

    default_df = next(iter(dataframes.values()))
    local_vars = {"df": default_df, "pd": pd}
    for name, frame in dataframes.items():
        safe = re.sub(r'[^a-zA-Z0-9_]', '_', name.rsplit('.', 1)[0])
        local_vars[safe] = frame

    try:
        exec(code, {}, local_vars)
        result = local_vars.get("result")
        if result is None:
            return {"code": code, "df": None, "summary": None,
                    "error": "No `result` variable produced.", "capped": False}
        if isinstance(result, pd.Series):
            result = result.reset_index()
            result.columns = [str(c) for c in result.columns]
        if not isinstance(result, pd.DataFrame):
            result = pd.DataFrame({"value": [str(result)]})
        capped = len(result) > MAX_DISPLAY_ROWS
        display_df = result.head(MAX_DISPLAY_ROWS)

        # Generate a natural language summary of the result
        summary = None
        try:
            summary_system = (
                "Summarize the following data result in 1-2 concise sentences "
                "that directly answer the user's question. Be specific with numbers."
            )
            summary_user = (
                f"User asked: {question}\n\n"
                f"Result ({len(result)} rows):\n{display_df.head(10).to_string(index=False)}"
            )
            summary = complete(summary_system, summary_user, 200)
        except Exception:
            pass  # summary is optional

        return {"code": code, "df": display_df, "summary": summary,
                "error": None, "capped": capped, "total_rows": len(result)}
    except Exception as e:
        return {"code": code, "df": None, "summary": None,
                "error": f"Execution error: {e}", "capped": False}


# ─────────────────────────────────────────────────────────────
# PIPELINE C — RAG  (with conversation memory + source citations)
# ─────────────────────────────────────────────────────────────

def run_rag_pipeline(complete, vectorstore, question: str, history: list) -> tuple[str, list[str]]:
    if vectorstore is None:
        return "No documents loaded. Please upload files in the sidebar first.", []

    docs = vectorstore.as_retriever(search_kwargs={"k": 8}).invoke(question)
    if not docs:
        return "I couldn't find relevant information in your uploaded documents.", []

    context = "\n\n---\n\n".join(
        f"[Source: {d.metadata.get('source', 'unknown')}"
        + (f", page {d.metadata['page']}" if 'page' in d.metadata else "")
        + f"]\n{d.page_content}"
        for d in docs
    )
    sources = list({d.metadata.get("source", "unknown") for d in docs})

    system = (
        "You are a helpful document assistant. Answer using ONLY the context below "
        "from the user's uploaded documents. Be concise but thorough. "
        "Cite sources by name when referencing specific information. "
        "If the answer isn't in the context, say so honestly.\n\n"
        f"Context:\n{context}"
    )
    # Include recent text-based conversation turns for follow-up questions
    recent_hist = [
        h for h in history[-MAX_HISTORY_TURNS*2:]
        if h["role"] in ("user", "assistant") and h.get("type") == "text"
    ]
    try:
        answer = complete(system, question, 1500, history=recent_hist)
    except Exception as e:
        return f"Error generating answer: {e}", sources

    return answer, sources


# ─────────────────────────────────────────────────────────────
# PIPELINE D — FORECAST  (predictive analytics)
# ─────────────────────────────────────────────────────────────

def _detect_date_col(df: pd.DataFrame) -> str | None:
    """Find first date-like column, trying parse if needed."""
    for col in df.columns:
        if pd.api.types.is_datetime64_any_dtype(df[col]):
            return col
        if df[col].dtype == object:
            try:
                parsed = pd.to_datetime(df[col], infer_datetime_format=True, errors="coerce")
                if parsed.notna().sum() / max(len(df), 1) > 0.7:
                    return col
            except Exception:
                pass
    return None


def _detect_numeric_cols(df: pd.DataFrame, exclude: list[str] | None = None) -> list[str]:
    exclude = exclude or []
    return [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c]) and c not in exclude]


def run_forecast_pipeline(
    complete,
    dataframes: dict,
    question: str,
    history: list,
) -> dict:
    """
    Predictive analytics pipeline:
    1. LLM selects which column to forecast + horizon + method
    2. sklearn linear / polynomial regression or moving-average extrapolation
    3. Prophet if installed (HAS_PROPHET)
    4. Returns Plotly figure + narrative summary
    """
    if not HAS_NUMPY or not HAS_SKLEARN:
        return {
            "fig": None, "summary": None,
            "error": (
                "scikit-learn / numpy not installed. "
                "Add `scikit-learn` and `numpy` to requirements.txt."
            ),
        }

    if not dataframes:
        return {"fig": None, "summary": None,
                "error": (
                    "📂 **No tabular data found.** The FORECAST pipeline requires a CSV or Excel file.\n\n"
                    "Please upload a `.csv`, `.xlsx`, or `.xls` file containing numeric columns "
                    "(e.g. sales over time, prices, quantities). PDF and Word documents cannot be forecasted — "
                    "they don't contain structured numeric series."
                )}

    # ── Step 1: LLM picks target col, x col, horizon, method ────
    df_summary = "\n".join(
        f"DataFrame '{name}': columns={list(df.columns)}, shape={df.shape}, "
        f"dtypes={df.dtypes.to_dict()}\nSample:\n{df.head(3).to_string(index=False)}"
        for name, df in dataframes.items()
    )
    plan_system = (
        "You are a forecasting assistant. Given the available DataFrames, decide:\n"
        "  - dataset: which dataset name to use\n"
        "  - x_col: column to use as the X axis (date/time preferred; null = use row index)\n"
        "  - y_col: numeric column to forecast\n"
        "  - horizon: number of future periods to predict (integer, default 10)\n"
        "  - method: 'linear', 'polynomial', or 'prophet' (use 'prophet' only if time-based)\n"
        "  - poly_degree: polynomial degree if method='polynomial' (2 or 3)\n"
        "  - title: short descriptive chart title\n\n"
        "Return ONLY a valid JSON object, no markdown fences, no explanation.\n\n"
        f"Available DataFrames:\n{df_summary}"
    )
    recent_hist = [
        h for h in history[-MAX_HISTORY_TURNS * 2:]
        if h["role"] in ("user", "assistant") and h.get("type") == "text"
    ]

    def _parse_plan_json(raw: str) -> dict:
        """Aggressively clean LLM output and parse as JSON, with multiple fallback strategies."""
        import json

        # Strip markdown fences (```json ... ``` or ``` ... ```)
        cleaned = re.sub(r"```[a-z]*\n?", "", raw.strip())
        cleaned = re.sub(r"```", "", cleaned).strip()

        # Extract first {...} block in case there's surrounding text
        brace_match = re.search(r"\{[\s\S]*\}", cleaned)
        if brace_match:
            cleaned = brace_match.group(0)

        # Fix common LLM JSON mistakes:
        # 1. Single quotes → double quotes (but not inside values)
        cleaned = re.sub(r"(?<![\\])\'", '"', cleaned)
        # 2. Trailing commas before } or ]
        cleaned = re.sub(r",\s*([}\]])", r"\1", cleaned)
        # 3. Python None/True/False → JSON null/true/false
        cleaned = cleaned.replace(": None", ": null").replace(":None", ": null")
        cleaned = cleaned.replace(": True", ": true").replace(":True", ": true")
        cleaned = cleaned.replace(": False", ": false").replace(":False", ": false")
        # 4. Unquoted null without colon
        cleaned = re.sub(r':\s*none\b', ': null', cleaned, flags=re.IGNORECASE)

        return json.loads(cleaned)

    plan = None
    try:
        raw = complete(plan_system, question, 400, history=recent_hist)
        plan = _parse_plan_json(raw)
    except Exception as e:
        # Full fallback: auto-detect best columns without LLM
        default_df = next(iter(dataframes.values()))
        num_cols = _detect_numeric_cols(default_df)
        date_col_auto = _detect_date_col(default_df)
        if not num_cols:
            return {"fig": None, "summary": None,
                    "error": f"Forecast planning failed and no numeric columns found. LLM error: {e}"}
        plan = {
            "dataset": next(iter(dataframes)),
            "x_col": date_col_auto,
            "y_col": num_cols[0],
            "horizon": 10,
            "method": "polynomial" if "polynomial" in question.lower() else "linear",
            "poly_degree": 2,
            "title": f"Forecast: {num_cols[0]}",
        }

    # ── Step 2: Resolve DataFrame ────────────────────────────────
    dataset_name = plan.get("dataset", next(iter(dataframes)))
    # fuzzy match
    df = dataframes.get(dataset_name) or next(iter(dataframes.values()))
    df = df.copy()

    y_col = plan.get("y_col")
    if not y_col or y_col not in df.columns:
        num_cols = _detect_numeric_cols(df)
        if not num_cols:
            return {"fig": None, "summary": None,
                    "error": "No numeric column found to forecast."}
        y_col = num_cols[0]

    x_col = plan.get("x_col")
    horizon = int(plan.get("horizon", FORECAST_HORIZON_DEFAULT))
    method = plan.get("method", "linear")
    poly_degree = int(plan.get("poly_degree", 2))
    title = plan.get("title", f"Forecast: {y_col}")

    # Determine x axis
    use_index = False
    date_col = None
    if x_col and x_col in df.columns:
        try:
            df[x_col] = pd.to_datetime(df[x_col], infer_datetime_format=True, errors="coerce")
            if df[x_col].isna().sum() / max(len(df), 1) < 0.3:
                date_col = x_col
            else:
                use_index = True
        except Exception:
            use_index = True
    else:
        detected = _detect_date_col(df)
        if detected:
            df[detected] = pd.to_datetime(df[detected], infer_datetime_format=True, errors="coerce")
            date_col = detected
        else:
            use_index = True

    # Drop rows where y is missing
    df = df.dropna(subset=[y_col])

    # ── Step 3: Build X numeric ──────────────────────────────────
    if date_col and not use_index:
        df = df.sort_values(date_col).reset_index(drop=True)
        df["_x_numeric"] = (df[date_col] - df[date_col].min()).dt.days.astype(float)
        x_label = date_col
    else:
        df = df.reset_index(drop=True)
        df["_x_numeric"] = df.index.astype(float)
        x_label = "Index"

    X = df["_x_numeric"].values.reshape(-1, 1)
    y = df[y_col].values.astype(float)

    # ── Step 4: Fit model ────────────────────────────────────────
    use_prophet = method == "prophet" and HAS_PROPHET and date_col
    r2 = None
    future_x_num = None
    future_x_dates = None

    if use_prophet:
        prophet_df = pd.DataFrame({"ds": df[date_col], "y": y})
        prophet_df = prophet_df.dropna()
        m = Prophet(yearly_seasonality="auto", weekly_seasonality="auto", daily_seasonality=False)
        m.fit(prophet_df)
        freq = pd.infer_freq(prophet_df["ds"]) or "D"
        future = m.make_future_dataframe(periods=horizon, freq=freq)
        forecast = m.predict(future)
        y_pred_hist = forecast[forecast["ds"].isin(prophet_df["ds"])]["yhat"].values
        r2 = r2_score(y, y_pred_hist[:len(y)])
        future_rows = forecast[~forecast["ds"].isin(prophet_df["ds"])].tail(horizon)
        future_x_dates = future_rows["ds"].values
        future_y_pred = future_rows["yhat"].values
        future_lower = future_rows["yhat_lower"].values
        future_upper = future_rows["yhat_upper"].values
    else:
        # sklearn: linear or polynomial
        if method == "polynomial":
            poly = PolynomialFeatures(degree=poly_degree, include_bias=False)
            X_poly = poly.fit_transform(X)
        else:
            X_poly = X
            poly = None
        reg = LinearRegression()
        reg.fit(X_poly, y)
        y_hist_pred = reg.predict(X_poly)
        r2 = r2_score(y, y_hist_pred)

        # Generate future X
        last_x = float(X[-1, 0])
        if date_col and not use_index:
            step = float(X[-1, 0] - X[-2, 0]) if len(X) > 1 else 1.0
            future_x_num = np.array([last_x + step * (i + 1) for i in range(horizon)])
            origin_date = df[date_col].min()
            future_x_dates = [origin_date + pd.Timedelta(days=int(d)) for d in future_x_num]
        else:
            step = 1.0
            future_x_num = np.array([last_x + step * (i + 1) for i in range(horizon)])

        fut_X = future_x_num.reshape(-1, 1)
        if poly:
            fut_X = poly.transform(fut_X)
        future_y_pred = reg.predict(fut_X)
        future_lower = future_upper = None  # no CI for sklearn without extra work

    # ── Step 5: Build Plotly figure ─────────────────────────────
    import plotly.graph_objects as go

    fig = go.Figure()

    # Historical data scatter
    if date_col and not use_index:
        hist_x = df[date_col]
    else:
        hist_x = df["_x_numeric"]

    fig.add_trace(go.Scatter(
        x=hist_x, y=y,
        mode="markers",
        name="Historical",
        marker=dict(color="#60a5fa", size=6, opacity=0.8),
    ))

    # Fitted line over history
    if use_prophet:
        hist_forecast = forecast[forecast["ds"].isin(df[date_col])]
        fig.add_trace(go.Scatter(
            x=hist_forecast["ds"], y=hist_forecast["yhat"],
            mode="lines",
            name="Fitted (Prophet)",
            line=dict(color="#818cf8", width=2),
        ))
    else:
        fit_x = hist_x
        if poly:
            y_fit_line = reg.predict(poly.transform(X))
        else:
            y_fit_line = reg.predict(X)
        fig.add_trace(go.Scatter(
            x=fit_x, y=y_fit_line,
            mode="lines",
            name=f"Fit ({method})",
            line=dict(color="#818cf8", width=2),
        ))

    # Forecast
    pred_x = future_x_dates if future_x_dates is not None else future_x_num
    fig.add_trace(go.Scatter(
        x=list(pred_x), y=future_y_pred,
        mode="lines+markers",
        name=f"Forecast (+{horizon})",
        line=dict(color="#4ade80", width=2, dash="dash"),
        marker=dict(color="#4ade80", size=7, symbol="diamond"),
    ))

    # Confidence band (Prophet only)
    if future_lower is not None and future_upper is not None:
        fig.add_trace(go.Scatter(
            x=list(pred_x) + list(pred_x)[::-1],
            y=list(future_upper) + list(future_lower)[::-1],
            fill="toself",
            fillcolor="rgba(74,222,128,0.12)",
            line=dict(color="rgba(0,0,0,0)"),
            name="95% CI",
            showlegend=True,
        ))

    # Divider between historical and forecast
    if len(pred_x) > 0:
        split_x = hist_x.iloc[-1] if hasattr(hist_x, "iloc") else hist_x[-1]
        fig.add_vline(
            x=split_x if not hasattr(split_x, "isoformat") else split_x,
            line_dash="dot", line_color="#4b5563",
            annotation_text="Forecast →",
            annotation_font_color="#6b7280",
            annotation_position="top right",
        )

    fig.update_layout(
        title=dict(text=title, font=dict(color="#818cf8", size=16)),
        paper_bgcolor="#0a0d14", plot_bgcolor="#111827",
        font_color="#e0e0e0", font_family="IBM Plex Sans",
        legend=dict(bgcolor="#111827", bordercolor="#1e2535"),
        xaxis=dict(title=x_label, gridcolor="#1e2535", zeroline=False),
        yaxis=dict(title=y_col, gridcolor="#1e2535", zeroline=False),
        hovermode="x unified",
    )

    # ── Step 6: Natural language summary ─────────────────────────
    summary = None
    try:
        last_actual = float(y[-1])
        last_forecast = float(future_y_pred[-1])
        direction = "↑ up" if last_forecast > last_actual else "↓ down"
        pct_change = abs((last_forecast - last_actual) / max(abs(last_actual), 1e-9)) * 100
        summary_prompt = (
            f"User asked: {question}\n\n"
            f"Forecast details:\n"
            f"- Column forecasted: {y_col}\n"
            f"- Method: {'Prophet' if use_prophet else method + ' regression'}\n"
            f"- R² (fit quality): {r2:.3f}\n"
            f"- Last historical value: {last_actual:.2f}\n"
            f"- Forecasted value at +{horizon} periods: {last_forecast:.2f} ({direction}, {pct_change:.1f}%)\n"
            f"- Horizon: {horizon} periods ahead\n\n"
            "Write a 2-3 sentence plain-English interpretation of these results. "
            "Be specific about the trend, confidence, and what the numbers mean."
        )
        summary = complete(
            "You are a data analyst. Summarize forecast results clearly and concisely.",
            summary_prompt, 300,
        )
    except Exception:
        pass

    return {
        "fig": fig,
        "summary": summary,
        "r2": r2,
        "method": "Prophet" if use_prophet else method,
        "horizon": horizon,
        "y_col": y_col,
        "error": None,
    }


# ─────────────────────────────────────────────────────────────
# PIPELINE E — SQL DATABASE
# ─────────────────────────────────────────────────────────────

DB_TYPE_DRIVERS = {
    "PostgreSQL":  {"driver": "postgresql+psycopg2", "port": 5432,  "pkg": "psycopg2-binary"},
    "MySQL":       {"driver": "mysql+pymysql",        "port": 3306,  "pkg": "pymysql"},
    "SQLite":      {"driver": "sqlite",               "port": None,  "pkg": None},
    "MSSQL":       {"driver": "mssql+pyodbc",         "port": 1433,  "pkg": "pyodbc"},
    "Oracle":      {"driver": "oracle+cx_oracle",     "port": 1521,  "pkg": "cx_oracle"},
}


def build_connection_url(db_type: str, host: str, port: int | None,
                          database: str, username: str, password: str,
                          sqlite_path: str = "") -> str:
    info = DB_TYPE_DRIVERS[db_type]
    driver = info["driver"]
    if db_type == "SQLite":
        return f"sqlite:///{sqlite_path}"
    if db_type == "MSSQL":
        return (f"mssql+pyodbc://{username}:{password}@{host}:{port}/{database}"
                "?driver=ODBC+Driver+17+for+SQL+Server")
    return f"{driver}://{username}:{password}@{host}:{port}/{database}"


def test_db_connection(url: str) -> tuple[bool, str, dict]:
    """Test connection and return (success, message, schema_info)."""
    if not HAS_SQLALCHEMY:
        return False, "SQLAlchemy not installed. Add `sqlalchemy` to requirements.txt.", {}
    try:
        engine = create_engine(url, connect_args={"connect_timeout": 8}, pool_pre_ping=True)
        with engine.connect() as conn:
            inspector = inspect(engine)
            tables = inspector.get_table_names()
            schema_info = {}
            for table in tables[:30]:  # cap at 30 tables
                try:
                    cols = inspector.get_columns(table)
                    schema_info[table] = [
                        {"name": c["name"], "type": str(c["type"])} for c in cols
                    ]
                except Exception:
                    schema_info[table] = []
        return True, f"Connected — {len(tables)} table(s) found", schema_info
    except Exception as e:
        return False, str(e), {}


def run_sql_query(url: str, sql: str, max_rows: int = MAX_DISPLAY_ROWS) -> tuple[pd.DataFrame | None, str | None]:
    """Execute a SQL query and return (DataFrame, error)."""
    try:
        engine = create_engine(url)
        with engine.connect() as conn:
            result = conn.execute(text(sql))
            rows = result.fetchmany(max_rows + 1)
            cols = list(result.keys())
            capped = len(rows) > max_rows
            df = pd.DataFrame(rows[:max_rows], columns=cols)
        return df, None
    except Exception as e:
        return None, str(e)


def nl_to_sql(complete, question: str, schema_info: dict, history: list, db_type: str) -> str:
    """Ask the LLM to generate SQL from natural language + schema."""
    schema_text = ""
    for table, cols in schema_info.items():
        col_defs = ", ".join(f"{c['name']} ({c['type']})" for c in cols)
        schema_text += f"  Table `{table}`: {col_defs}\n"

    dialect_hints = {
        "PostgreSQL": "Use PostgreSQL syntax. Use LIMIT for row limits.",
        "MySQL":      "Use MySQL syntax. Use LIMIT for row limits.",
        "SQLite":     "Use SQLite syntax. Use LIMIT for row limits.",
        "MSSQL":      "Use T-SQL (SQL Server) syntax. Use TOP N for row limits.",
        "Oracle":     "Use Oracle SQL syntax. Use ROWNUM or FETCH FIRST N ROWS for row limits.",
    }.get(db_type, "Use standard SQL.")

    system = (
        f"You are an expert SQL assistant. {dialect_hints}\n\n"
        f"Database schema:\n{schema_text}\n"
        "Generate a single valid SQL query that answers the user's question.\n"
        f"Return ONLY the raw SQL query, no markdown fences, no explanation, no semicolon at end."
    )
    recent_hist = [
        h for h in history[-MAX_HISTORY_TURNS * 2:]
        if h["role"] in ("user", "assistant") and h.get("type") == "text"
    ]
    raw = complete(system, question, 600, history=recent_hist)
    # Strip fences
    raw = re.sub(r"^```[a-z]*\n?", "", raw.strip(), flags=re.MULTILINE)
    raw = re.sub(r"```$", "", raw.strip(), flags=re.MULTILINE)
    raw = raw.strip().rstrip(";")
    return raw


# ─────────────────────────────────────────────────────────────
# EXAMPLE PROMPTS
# ─────────────────────────────────────────────────────────────

EXAMPLE_PROMPTS = [
    ("📊", "Plot a bar chart of the most frequent values in the first column"),
    ("📊", "Show a line chart of numeric columns over time"),
    ("📊", "Visualize the distribution with a histogram"),
    ("🔢", "Show me the first 20 rows"),
    ("🔢", "What are the column names and data types?"),
    ("🔢", "Show rows where any value is missing"),
    ("🔢", "Sort by the first numeric column, descending, top 10"),
    ("📋", "Summarize the main topics in the documents"),
    ("📋", "What are the key findings or conclusions?"),
    ("📋", "List all important names, dates, or numbers mentioned"),
    ("🔮", "Forecast the next 10 periods using linear regression"),
    ("🔮", "Predict future values and show trend line"),
    ("🔮", "Run a polynomial regression and extrapolate"),
    ("🗄️", "Show me all tables and their row counts"),
    ("🗄️", "What are the top 10 records by the first numeric column?"),
    ("🗄️", "Show me a summary of the database schema"),
]


# ─────────────────────────────────────────────────────────────
# SIDEBAR STATS HELPER
# ─────────────────────────────────────────────────────────────

def render_sidebar_stats(uploaded_files, dataframes):
    if not uploaded_files:
        return
    st.markdown("---")
    st.markdown("### 📁 Loaded Files")
    for f in uploaded_files:
        ext = f.name.rsplit(".", 1)[-1].upper()
        color = {"PDF":"#f59e0b","DOCX":"#3b82f6","CSV":"#10b981",
                 "XLSX":"#10b981","XLS":"#10b981","TXT":"#8b5cf6"}.get(ext, "#6b7280")
        st.markdown(
            f"<div style='font-size:12px; padding:4px 0; color:{color}; "
            f"border-left:2px solid {color}; padding-left:8px; margin:3px 0;'>"
            f"<b>{ext}</b> {f.name}</div>",
            unsafe_allow_html=True,
        )
    if dataframes:
        st.markdown("---")
        st.markdown("### 📊 Tabular Data")
        for name, df in dataframes.items():
            st.markdown(
                f"<div class='stat-card' style='margin:4px 0;'>"
                f"<div style='font-size:11px; color:#9ca3af;'>{name}</div>"
                f"<div style='font-size:13px; color:#60a5fa; font-weight:600;'>"
                f"{df.shape[0]:,} rows · {df.shape[1]} cols</div></div>",
                unsafe_allow_html=True,
            )


# ─────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────

def main():
    # ── Session state init ────────────────────────────────
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "selected_model" not in st.session_state:
        st.session_state.selected_model = DEFAULT_MODEL

    all_model_ids = list(MODELS.keys())

    # ── Sidebar ───────────────────────────────────────────
    with st.sidebar:
        st.markdown("## 🤖 DocBot")
        st.markdown("*AI-powered document assistant*")
        st.markdown("---")

        st.markdown("### 🤖 Model")
        selected_model = st.selectbox(
            "Model",
            options=all_model_ids,
            index=all_model_ids.index(
                st.session_state.selected_model
                if st.session_state.selected_model in all_model_ids
                else DEFAULT_MODEL
            ),
            format_func=lambda mid: f"{MODELS[mid]['icon']}  {MODELS[mid]['label']}",
            label_visibility="collapsed",
        )
        if selected_model != st.session_state.selected_model:
            st.session_state.selected_model = selected_model
            # Clear client cache for old model implicitly via new key

        m = MODELS[selected_model]
        st.markdown(
            f"<div class='model-card' style='border-color:{m['color']}33; background:#111827;'>"
            f"<b style='color:{m['color']}'>{m['icon']} {m['label']}</b><br>"
            f"<span style='color:#6b7280; font-size:11px'>{m['description']}</span></div>",
            unsafe_allow_html=True,
        )

        st.markdown("---")
        st.markdown("""
        <div class='info-box'>
        <b style='color:#c084fc'>📊 GRAPH</b> — Charts from CSV / Excel<br>
        <b style='color:#60a5fa'>🔵 DATA</b> &nbsp;— Query & filter rows<br>
        <b style='color:#4ade80'>🟢 RAG</b> &nbsp;&nbsp;— Q&A from documents<br>
        <b style='color:#f59e0b'>🔮 FORECAST</b> — Predict & extrapolate<br>
        <b style='color:#38bdf8'>🗄️ SQL</b> &nbsp;&nbsp;&nbsp;— Query your database<br><br>
        <span style='color:#6b7280; font-size:11px'>Follow-up questions remember previous answers</span>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🗑️ Clear Chat", use_container_width=True):
                st.session_state.messages = []
                st.rerun()
        with col2:
            msg_count = len([m for m in st.session_state.messages if m["role"] == "user"])
            st.markdown(
                f"<div style='text-align:center; font-size:11px; color:#6b7280; padding-top:8px'>"
                f"{msg_count} turn{'s' if msg_count != 1 else ''}</div>",
                unsafe_allow_html=True,
            )

        # ── SQL Database Panel ───────────────────────────────
        st.markdown("---")
        st.markdown("### 🗄️ SQL Database")

        # Init SQL session state
        if "sql_connected" not in st.session_state:
            st.session_state.sql_connected = False
        if "sql_conn_url" not in st.session_state:
            st.session_state.sql_conn_url = ""
        if "sql_schema" not in st.session_state:
            st.session_state.sql_schema = {}
        if "sql_db_type" not in st.session_state:
            st.session_state.sql_db_type = "PostgreSQL"

        if st.session_state.sql_connected:
            n_tables = len(st.session_state.sql_schema)
            st.markdown(
                f"<div class='sql-connected'>✅ Connected · {n_tables} table(s)</div>",
                unsafe_allow_html=True,
            )
            chips_html = "<div style='margin:6px 0 4px 0'>"
            for tbl in list(st.session_state.sql_schema.keys())[:12]:
                chips_html += f"<span class='sql-table-chip'>⬡ {tbl}</span>"
            if n_tables > 12:
                chips_html += f"<span style='color:#4b5563; font-size:11px'> +{n_tables-12} more</span>"
            chips_html += "</div>"
            st.markdown(chips_html, unsafe_allow_html=True)
            if st.button("🔌 Disconnect", use_container_width=True):
                st.session_state.sql_connected = False
                st.session_state.sql_conn_url = ""
                st.session_state.sql_schema = {}
                st.rerun()
        else:
            if not HAS_SQLALCHEMY:
                st.warning("Install `sqlalchemy` to enable SQL integration.")
            else:
                with st.form("sql_connect_form"):
                    db_type = st.selectbox(
                        "Database type",
                        options=list(DB_TYPE_DRIVERS.keys()),
                    )
                    if db_type == "SQLite":
                        f_sqlite = st.text_input("SQLite file path", placeholder="/path/to/db.sqlite")
                        f_host = f_db = f_user = f_pass = ""
                        f_port = 0
                    else:
                        col_h, col_p = st.columns([3, 1])
                        with col_h:
                            f_host = st.text_input("Host", placeholder="localhost")
                        with col_p:
                            f_port = st.number_input("Port", value=DB_TYPE_DRIVERS[db_type]["port"], step=1)
                        f_db   = st.text_input("Database name")
                        f_user = st.text_input("Username")
                        f_pass = st.text_input("Password", type="password")
                        f_sqlite = ""
                    submitted = st.form_submit_button("🔗 Test & Connect", use_container_width=True, type="primary")

                if submitted:
                    st.session_state.sql_db_type = db_type
                    if db_type == "SQLite":
                        if not f_sqlite.strip():
                            st.error("Please enter the SQLite file path.")
                            st.stop()
                        _url = f"sqlite:///{f_sqlite.strip()}"
                    else:
                        missing = [l for l, v in [("Host", f_host), ("Database name", f_db), ("Username", f_user)] if not v.strip()]
                        if missing:
                            st.error(f"Please fill in: {', '.join(missing)}")
                            st.stop()
                        _url = build_connection_url(db_type, f_host.strip(), int(f_port), f_db.strip(), f_user.strip(), f_pass)
                    with st.spinner("Connecting…"):
                        ok, msg, schema = test_db_connection(_url)
                    if ok:
                        st.session_state.sql_connected = True
                        st.session_state.sql_conn_url = _url
                        st.session_state.sql_schema = schema
                        st.rerun()
                    else:
                        st.error(f"❌ {msg}")

        # File stats in sidebar
        uploaded_files_sidebar = st.session_state.get("_uploaded_files_meta", [])

    # ── File uploader — main area ─────────────────────────
    st.markdown("## 🤖 DocBot")

    _, upload_col, _ = st.columns([1, 2, 1])
    with upload_col:
        uploaded_files = st.file_uploader(
            "📁 Drop files here or click to browse",
            type=["pdf", "docx", "txt", "csv", "xlsx", "xls"],
            accept_multiple_files=True,
            label_visibility="visible",
        )

    vectorstore = None
    dataframes = {}

    if uploaded_files:
        # File chips
        file_chips_html = "<div style='text-align:center; margin: 6px 0 12px 0;'>"
        for f in uploaded_files:
            ext = f.name.rsplit(".", 1)[-1].upper()
            color = {"PDF":"#f59e0b","DOCX":"#3b82f6","CSV":"#10b981",
                     "XLSX":"#10b981","XLS":"#10b981","TXT":"#8b5cf6"}.get(ext, "#6b7280")
            file_chips_html += (
                f"<span style='background:#111827; border:1px solid {color}44; color:{color};"
                f"border-radius:20px; padding:3px 12px; font-size:12px; display:inline-block;"
                f"margin:3px; font-family:IBM Plex Mono,monospace'>{ext} · {f.name}</span>"
            )
        file_chips_html += "</div>"
        st.markdown(file_chips_html, unsafe_allow_html=True)

        file_contents = tuple((f.name, f.read()) for f in uploaded_files)
        for f in uploaded_files:
            f.seek(0)

        # Use stable content hash as cache key
        chash = _content_hash(file_contents)
        vectorstore, dataframes = build_vectorstore(chash, file_contents)

        # Store metadata for sidebar
        st.session_state["_uploaded_files_meta"] = [(f.name,) for f in uploaded_files]
        st.session_state["_dataframes_meta"] = {
            name: df.shape for name, df in dataframes.items()
        }

        # Render file stats in sidebar
        with st.sidebar:
            st.markdown("---")
            st.markdown("### 📁 Loaded")
            for f in uploaded_files:
                ext = f.name.rsplit(".", 1)[-1].upper()
                color = {"PDF":"#f59e0b","DOCX":"#3b82f6","CSV":"#10b981",
                         "XLSX":"#10b981","XLS":"#10b981","TXT":"#8b5cf6"}.get(ext, "#6b7280")
                st.markdown(
                    f"<div style='font-size:12px; padding:3px 0; color:{color}; "
                    f"border-left:2px solid {color}44; padding-left:8px; margin:2px 0;'>"
                    f"{f.name}</div>",
                    unsafe_allow_html=True,
                )
            if dataframes:
                st.markdown("---")
                for name, df in dataframes.items():
                    st.markdown(
                        f"<div class='stat-card' style='margin:4px 0;'>"
                        f"<div style='font-size:10px; color:#6b7280; margin-bottom:2px'>{name}</div>"
                        f"<div style='font-size:13px; color:#60a5fa; font-weight:600;'>"
                        f"{df.shape[0]:,} rows · {df.shape[1]} cols</div></div>",
                        unsafe_allow_html=True,
                    )

    st.markdown("---")

    # ── No files + no DB guard ────────────────────────────
    sql_connected = st.session_state.get("sql_connected", False)
    if not uploaded_files and not sql_connected:
        st.markdown("""
        <div class='upload-hint'>
            <p style='font-size:16px; color:#818cf8; font-weight:700; margin-bottom:8px'>
                Upload documents or connect a database to get started
            </p>
            <p style='font-size:13px; color:#6b7280; line-height:1.8'>
                PDF · Word · CSV · Excel · TXT<br>
                Summarize documents · Plot charts · Filter rows · Answer questions<br><br>
                <span style='color:#38bdf8'>🗄️ Connect a SQL database in the sidebar to query it with plain English</span>
            </p>
        </div>
        """, unsafe_allow_html=True)
        st.stop()

    # ── Build LLM client (with auto-fallback) ────────────
    model_id = st.session_state.selected_model
    try:
        _get_client(model_id)  # validate primary credentials eagerly
        complete = make_fallback_complete(model_id)
    except Exception as e:
        st.error(
            f"❌ Could not initialise model **{model_id}**.\n\n"
            f"**Error:** {e}\n\n"
            "Check your `.env` credentials — see model description for required keys."
        )
        st.stop()

    # ── Sub-header ────────────────────────────────────────
    m = MODELS[model_id]
    active_label = m["label"]
    active_color = m["color"]
    active_icon  = m["icon"]
    # Show which model actually handled last request (if fallback occurred)
    last_fallback = st.session_state.get("_last_fallback")
    if last_fallback and last_fallback in MODELS:
        fm = MODELS[last_fallback]
        active_label = f"{m['label']} <span style='color:#f59e0b; font-size:11px'>→ fell back to {fm['icon']} {fm['label']}</span>"

    parts = [
        f"<span style='color:{active_color}'>{active_icon} {active_label}</span>",
    ]
    if uploaded_files:
        parts.append(f"📄 {len(uploaded_files)} file(s)")
    if dataframes:
        parts.append(f"📊 {len(dataframes)} tabular dataset(s)")
    if sql_connected:
        n_tbls = len(st.session_state.get("sql_schema", {}))
        parts.append(f"<span style='color:#38bdf8'>🗄️ DB · {n_tbls} tables</span>")
    turn_count = len([msg for msg in st.session_state.messages if msg["role"] == "user"])
    if turn_count > 0:
        parts.append(f"💬 {turn_count} turn(s)")
    st.markdown(
        f"<p style='color:#4b5563; margin-top:-8px; font-size:13px'>"
        + " &nbsp;·&nbsp; ".join(parts) + "</p>",
        unsafe_allow_html=True,
    )

    # ── Chat history ──────────────────────────────────────
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            if msg.get("type") == "plotly":
                st.plotly_chart(msg["content"], use_container_width=True)
            elif msg.get("type") == "forecast":
                c = msg["content"]
                if c.get("summary"):
                    st.markdown(c["summary"])
                if c.get("meta"):
                    st.markdown(c["meta"], unsafe_allow_html=True)
                if c.get("fig"):
                    st.plotly_chart(c["fig"], use_container_width=True)
            elif msg.get("type") == "dataframe":
                c = msg["content"]
                if c.get("summary"):
                    st.markdown(c["summary"])
                st.markdown(c["header"])
                if c.get("warning"):
                    st.warning(c["warning"])
                if c.get("df") is not None and not c["df"].empty:
                    st.dataframe(c["df"], use_container_width=True, hide_index=True)
                    # Download button (CSV)
                    csv_bytes = c["df"].to_csv(index=False).encode()
                    st.download_button(
                        "⬇️ Download CSV",
                        data=csv_bytes,
                        file_name="result.csv",
                        mime="text/csv",
                        key=f"dl_{id(c)}",
                    )
            else:
                st.markdown(msg["content"])
                # Show source chips for RAG answers
                if msg.get("sources"):
                    chips = "".join(
                        f"<span class='source-chip'>📄 {s}</span>"
                        for s in msg["sources"]
                    )
                    st.markdown(
                        f"<div style='margin-top:6px'>{chips}</div>",
                        unsafe_allow_html=True,
                    )
            if msg.get("timing"):
                st.markdown(
                    f"<div class='timing-note'>⏱ {msg['timing']:.1f}s</div>",
                    unsafe_allow_html=True,
                )

    # ── Example prompt chips — shown only when no messages ─
    if not st.session_state.messages:
        chip_groups = [
            ("📊", "#2d1a4a", "#c084fc", [p for e, p in EXAMPLE_PROMPTS if e == "📊"]),
            ("🔢", "#1a2e4a", "#60a5fa", [p for e, p in EXAMPLE_PROMPTS if e == "🔢"]),
            ("📋", "#1a3a2a", "#4ade80", [p for e, p in EXAMPLE_PROMPTS if e == "📋"]),
            ("🔮", "#2a1e10", "#f59e0b", [p for e, p in EXAMPLE_PROMPTS if e == "🔮"]),
        ]
        if st.session_state.get("sql_connected"):
            chip_groups.append(("🗄️", "#0d1829", "#38bdf8", [p for e, p in EXAMPLE_PROMPTS if e == "🗄️"]))
        st.markdown(
            "<p style='color:#6b7280; font-size:12px; margin-bottom:6px; text-align:center'>"
            "💡 Try one of these prompts</p>",
            unsafe_allow_html=True,
        )
        for emoji, bg, fg, prompts in chip_groups:
            cols = st.columns(len(prompts))
            for col, prompt in zip(cols, prompts):
                with col:
                    short = prompt if len(prompt) <= 38 else prompt[:35] + "…"
                    if st.button(f"{emoji} {short}", key=f"chip_{prompt}", use_container_width=True):
                        st.session_state.pending_question = prompt
                        st.rerun()
        st.markdown("<div style='margin-bottom:8px'></div>", unsafe_allow_html=True)

    # ── Chat input ────────────────────────────────────────
    if "pending_question" in st.session_state:
        user_input = st.session_state.pop("pending_question")
    else:
        user_input = st.chat_input("Ask about your documents or database…")

    if user_input:
        st.session_state.messages.append({"role": "user", "content": user_input, "type": "text"})
        with st.chat_message("user"):
            st.markdown(user_input)

        # Build conversation history for LLM context
        history = [
            {"role": m["role"], "content": m["content"], "type": m.get("type", "text")}
            for m in st.session_state.messages[:-1]  # exclude the just-added user msg
            if m.get("type") == "text" and isinstance(m.get("content"), str)
        ]

        with st.chat_message("assistant"):
            t_start = time.time()
            with st.spinner("🤔 Classifying…"):
                intents = classify_intents(
                    complete, user_input,
                    has_tabular=bool(dataframes),
                    has_sql=st.session_state.get("sql_connected", False),
                )

            # Show all intent badges upfront so user knows what's coming
            badge_html = ""
            for intent in intents:
                if intent == "GRAPH":
                    badge_html += '<span class="badge-graph">📊 GRAPH</span> '
                elif intent == "DATA":
                    badge_html += '<span class="badge-data">🔵 DATA</span> '
                elif intent == "FORECAST":
                    badge_html += '<span class="badge-forecast">🔮 FORECAST</span> '
                elif intent == "SQL":
                    badge_html += '<span class="badge-sql">🗄️ SQL</span> '
                else:
                    badge_html += '<span class="badge-rag">🟢 RAG</span> '
            st.markdown(badge_html, unsafe_allow_html=True)

            # ── Execute each intent in order ───────────────
            for intent in intents:

                # ── RAG ────────────────────────────────────
                if intent == "RAG":
                    with st.spinner("📚 Searching documents…"):
                        t0 = time.time()
                        answer, sources = run_rag_pipeline(complete, vectorstore, user_input, history)
                        elapsed = time.time() - t0
                    st.markdown(answer)
                    if sources:
                        chips = "".join(
                            f"<span class='source-chip'>📄 {s}</span>" for s in sources
                        )
                        st.markdown(
                            f"<div style='margin-top:8px'>{chips}</div>",
                            unsafe_allow_html=True,
                        )
                    st.markdown(
                        f"<div class='timing-note'>⏱ {elapsed:.1f}s</div>",
                        unsafe_allow_html=True,
                    )
                    st.session_state.messages.append({
                        "role": "assistant", "content": answer, "type": "text",
                        "sources": sources, "timing": elapsed,
                    })

                # ── GRAPH ──────────────────────────────────
                elif intent == "GRAPH":
                    with st.spinner("🎨 Generating chart…"):
                        t0 = time.time()
                        fig, err = run_graph_pipeline(complete, dataframes, user_input, history)
                        # No tabular data — extract numbers from document text instead
                        if err == "NO_TABULAR":
                            with st.spinner("📊 Extracting data from documents…"):
                                fig, err = run_graph_from_context(complete, vectorstore, user_input, history)
                        elapsed = time.time() - t0
                    if fig:
                        st.plotly_chart(fig, use_container_width=True)
                        st.markdown(
                            f"<div class='timing-note'>⏱ {elapsed:.1f}s</div>",
                            unsafe_allow_html=True,
                        )
                        st.session_state.messages.append({
                            "role": "assistant", "content": fig,
                            "type": "plotly", "timing": elapsed,
                        })
                    else:
                        st.markdown('<span class="badge-error">❌ CHART ERROR</span>', unsafe_allow_html=True)
                        st.error(err)
                        st.session_state.messages.append({
                            "role": "assistant", "content": err, "type": "text",
                        })

                # ── DATA ───────────────────────────────────
                elif intent == "DATA":
                    with st.spinner("🔍 Querying data…"):
                        t0 = time.time()
                        result = run_data_pipeline(complete, dataframes, user_input, history)
                        elapsed = time.time() - t0

                    if result["error"]:
                        st.markdown('<span class="badge-error">❌ DATA ERROR</span>', unsafe_allow_html=True)
                        st.error(result["error"])
                        st.session_state.messages.append({
                            "role": "assistant", "content": result["error"], "type": "text",
                        })
                    else:
                        df = result["df"]
                        summary = result.get("summary")
                        if summary:
                            st.markdown(summary)
                        total = result.get("total_rows", len(df) if df is not None else 0)
                        header = (
                            f"<details><summary style='cursor:pointer; color:#6b7280; "
                            f"font-size:12px; font-family:IBM Plex Mono,monospace'>🔧 Code</summary>"
                            f"\n\n```python\n{result['code']}\n```\n</details>"
                        )
                        warning = None
                        if df is not None and not df.empty:
                            if result["capped"]:
                                warning = (
                                    f"⚠️ Showing first {MAX_DISPLAY_ROWS} of {total:,} rows. "
                                    "Refine your question to narrow results."
                                )
                            row_label = f"**{total:,} row{'s' if total != 1 else ''}** returned"
                            st.markdown(row_label)
                            if warning:
                                st.warning(warning)
                            st.markdown(header, unsafe_allow_html=True)
                            st.dataframe(df, use_container_width=True, hide_index=True)
                            csv_bytes = df.to_csv(index=False).encode()
                            st.download_button(
                                "⬇️ Download CSV",
                                data=csv_bytes,
                                file_name="result.csv",
                                mime="text/csv",
                                key=f"dl_new_{len(st.session_state.messages)}_{intent}",
                            )
                        else:
                            st.markdown("No results found.")
                        st.markdown(
                            f"<div class='timing-note'>⏱ {elapsed:.1f}s</div>",
                            unsafe_allow_html=True,
                        )
                        st.session_state.messages.append({
                            "role": "assistant", "type": "dataframe", "timing": elapsed,
                            "content": {
                                "header": header, "warning": warning, "df": df,
                                "summary": summary,
                            },
                        })

                # ── FORECAST ───────────────────────────────
                elif intent == "FORECAST":
                    with st.spinner("🔮 Running predictive model…"):
                        t0 = time.time()
                        fc = run_forecast_pipeline(complete, dataframes, user_input, history)
                        elapsed = time.time() - t0

                    if fc["error"]:
                        st.markdown('<span class="badge-error">❌ FORECAST ERROR</span>', unsafe_allow_html=True)
                        st.error(fc["error"])
                        st.session_state.messages.append({
                            "role": "assistant", "content": fc["error"], "type": "text",
                        })
                    else:
                        if fc.get("summary"):
                            st.markdown(fc["summary"])
                        # Metadata pill row
                        r2_color = "#4ade80" if (fc.get("r2") or 0) > 0.8 else "#f59e0b" if (fc.get("r2") or 0) > 0.5 else "#f87171"
                        meta_html = (
                            f"<div style='margin: 6px 0 10px 0; font-size:12px; font-family:IBM Plex Mono,monospace'>"
                            f"<span style='background:#1a2e4a; color:#60a5fa; padding:2px 8px; border-radius:4px; margin-right:6px'>"
                            f"method: {fc.get('method','?')}</span>"
                            f"<span style='background:#1a3a2a; color:{r2_color}; padding:2px 8px; border-radius:4px; margin-right:6px'>"
                            f"R²: {fc.get('r2', 0):.3f}</span>"
                            f"<span style='background:#2a1e10; color:#f59e0b; padding:2px 8px; border-radius:4px; margin-right:6px'>"
                            f"horizon: +{fc.get('horizon','?')} periods</span>"
                            f"<span style='background:#1e1a2a; color:#c084fc; padding:2px 8px; border-radius:4px'>"
                            f"target: {fc.get('y_col','?')}</span>"
                            f"</div>"
                        )
                        st.markdown(meta_html, unsafe_allow_html=True)
                        if fc.get("fig"):
                            st.plotly_chart(fc["fig"], use_container_width=True)
                        st.markdown(
                            f"<div class='timing-note'>⏱ {elapsed:.1f}s</div>",
                            unsafe_allow_html=True,
                        )
                        st.session_state.messages.append({
                            "role": "assistant", "type": "forecast", "timing": elapsed,
                            "content": {
                                "fig": fc["fig"],
                                "summary": fc.get("summary"),
                                "meta": meta_html,
                            },
                        })

                # ── SQL ────────────────────────────────────
                elif intent == "SQL":
                    if not st.session_state.get("sql_connected"):
                        st.markdown('<span class="badge-error">❌ SQL ERROR</span>', unsafe_allow_html=True)
                        err_msg = "No database connected. Use the **🗄️ SQL Database** panel in the sidebar to connect."
                        st.error(err_msg)
                        st.session_state.messages.append({"role": "assistant", "content": err_msg, "type": "text"})
                    else:
                        with st.spinner("🗄️ Generating SQL query…"):
                            t0 = time.time()
                            conn_url = st.session_state.sql_conn_url
                            schema = st.session_state.sql_schema
                            db_type = st.session_state.sql_db_type
                            sql_query = nl_to_sql(complete, user_input, schema, history, db_type)

                        with st.spinner("⚡ Running query…"):
                            result_df, sql_err = run_sql_query(conn_url, sql_query)
                            elapsed = time.time() - t0

                        # Always show the generated SQL in a collapsible block
                        sql_block_html = (
                            f"<details><summary style='cursor:pointer; color:#38bdf8; "
                            f"font-size:12px; font-family:IBM Plex Mono,monospace'>🗄️ Generated SQL</summary>"
                            f"\n\n```sql\n{sql_query}\n```\n</details>"
                        )

                        if sql_err:
                            st.markdown('<span class="badge-error">❌ SQL ERROR</span>', unsafe_allow_html=True)
                            st.markdown(sql_block_html, unsafe_allow_html=True)
                            # Try to self-heal: ask LLM to fix the SQL
                            with st.spinner("🔧 Attempting to fix query…"):
                                fix_system = (
                                    f"The following SQL query produced an error. Fix it.\n"
                                    f"Error: {sql_err}\n"
                                    f"Original query:\n{sql_query}\n\n"
                                    f"Database schema:\n"
                                    + "\n".join(
                                        f"  Table `{t}`: " + ", ".join(f"{c['name']} ({c['type']})" for c in cols)
                                        for t, cols in schema.items()
                                    )
                                    + "\n\nReturn ONLY the corrected SQL, no explanation, no fences."
                                )
                                try:
                                    fixed_sql = complete(fix_system, "Fix the SQL", 400)
                                    fixed_sql = re.sub(r"^```[a-z]*\n?", "", fixed_sql.strip(), flags=re.MULTILINE)
                                    fixed_sql = re.sub(r"```$", "", fixed_sql.strip()).strip().rstrip(";")
                                    result_df, sql_err2 = run_sql_query(conn_url, fixed_sql)
                                    if sql_err2 is None:
                                        sql_query = fixed_sql
                                        sql_block_html = (
                                            f"<details open><summary style='cursor:pointer; color:#4ade80; "
                                            f"font-size:12px; font-family:IBM Plex Mono,monospace'>✅ Fixed SQL</summary>"
                                            f"\n\n```sql\n{sql_query}\n```\n</details>"
                                        )
                                        st.markdown(sql_block_html, unsafe_allow_html=True)
                                        sql_err = None
                                    else:
                                        st.error(f"Query failed: {sql_err}\n\nAuto-fix also failed: {sql_err2}")
                                        st.session_state.messages.append({
                                            "role": "assistant", "content": f"SQL error: {sql_err}", "type": "text"
                                        })
                                except Exception as fix_e:
                                    st.error(f"Query failed: {sql_err}")
                                    st.session_state.messages.append({
                                        "role": "assistant", "content": f"SQL error: {sql_err}", "type": "text"
                                    })

                        if sql_err is None and result_df is not None:
                            # NL summary of results
                            summary = None
                            try:
                                summary_prompt = (
                                    f"User asked: {user_input}\n\n"
                                    f"SQL executed:\n{sql_query}\n\n"
                                    f"Result ({len(result_df)} rows):\n{result_df.head(10).to_string(index=False)}\n\n"
                                    "Summarize the result in 1-2 sentences that directly answer the question."
                                )
                                summary = complete(
                                    "You are a data analyst. Be concise and specific with numbers.",
                                    summary_prompt, 200
                                )
                            except Exception:
                                pass

                            if summary:
                                st.markdown(summary)
                            st.markdown(sql_block_html, unsafe_allow_html=True)

                            capped = len(result_df) >= MAX_DISPLAY_ROWS
                            row_label = f"**{len(result_df):,} row{'s' if len(result_df) != 1 else ''}** returned"
                            if capped:
                                row_label += f" *(showing first {MAX_DISPLAY_ROWS})*"
                            st.markdown(row_label)
                            st.dataframe(result_df, use_container_width=True, hide_index=True)
                            csv_bytes = result_df.to_csv(index=False).encode()
                            st.download_button(
                                "⬇️ Download CSV",
                                data=csv_bytes,
                                file_name="query_result.csv",
                                mime="text/csv",
                                key=f"dl_sql_{len(st.session_state.messages)}",
                            )
                            st.markdown(
                                f"<div class='timing-note'>⏱ {elapsed:.1f}s</div>",
                                unsafe_allow_html=True,
                            )
                            st.session_state.messages.append({
                                "role": "assistant", "type": "dataframe", "timing": elapsed,
                                "content": {
                                    "header": sql_block_html,
                                    "warning": f"Showing first {MAX_DISPLAY_ROWS} rows." if capped else None,
                                    "df": result_df,
                                    "summary": summary,
                                },
                            })

                # Divider between multiple pipeline outputs
                if len(intents) > 1 and intent != intents[-1]:
                    st.markdown("<hr style='border-color:#1e2535; margin:16px 0'>", unsafe_allow_html=True)


if __name__ == "__main__":
    main()